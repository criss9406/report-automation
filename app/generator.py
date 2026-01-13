#generator.py

from os import path
from docx import Document
from datetime import datetime
import polars as pl
from pathlib import Path
from app.logger import configurar_logger

logger = configurar_logger('generator')


def crear_reporte(df: pl.dataframe) -> str:

    """
    Genera reporte DOCX a partir de DataFrame procesado.
    
    Args:
        df: DataFrame con datos de población mundial
        
    Returns:
        str: Ruta completa del archivo generado
        None: Si hay error crítico
    """
    try:
        logger.info("=" * 60)
        logger.info("INICIANDO GENERACIÓN DE REPORTE")
        logger.info("=" * 60)

        #validar entrada
        if df is None or df.height == 0:
            logger.error("❌ DF está vacío o es None")
            return None

        logger.info(f"  DF recibido: {df.height} registros")

        logger.info("  verificando existencia de carpeta outputs")
        Path("outputs").mkdir(exist_ok=True)

        logger.info("   creando documento word")
        doc = Document()

        #titulo principal
        titulo = doc.add_heading("reporte de la pobración mundial", 0)

        #fecha
        fecha_actual = datetime.now().strftime("%d/%m/%Y %H:%M")
        doc.add_paragraph(f"fecha de generación: {fecha_actual}")
        doc.add_paragraph()  #espacio en blanco

        logger.info("   calculando estadísticas globales...")

        #estadísticas globales
        doc.add_heading("Estadísticas globales", 1)

        #calcular estadísticas
        poblacion_total = df["poblacion_2024"].sum()
        poblacion_promedio = df["poblacion_2024"].mean()
        total_paises = df.height

        logger.info(f"  - población mundial total: {poblacion_total:,}")
        logger.info(f"  - promedio por país: {int(poblacion_promedio):,}")
        logger.info(f"  - total países: {total_paises}")

        #crear tabla de estadísticas
        tabla_stats = doc.add_table(rows=4, cols=2)
        tabla_stats.style = 'Light Grid Accent 1'

        #encabezados
        tabla_stats.rows[0].cells[0].text = "métrica"
        tabla_stats.rows[0].cells[1].text = "valor"

        #datos
        tabla_stats.rows[1].cells[0].text = "Población mundial"
        tabla_stats.rows[1].cells[1].text = f"{poblacion_total:,}"

        tabla_stats.rows[2].cells[0].text = "población promedio por país"
        tabla_stats.rows[2].cells[1].text = f"{int(poblacion_promedio):,}"

        tabla_stats.rows[3].cells[0].text = "total de países"
        tabla_stats.rows[3].cells[1].text = str(total_paises)

        doc.add_paragraph()

        logger.info("   generando tabla TOP 10...")
        #top 10 países
        doc.add_heading("top 10 países por población", 1)

        top10 = df.head(10)

        #crear tabla
        tabla_top = doc.add_table(rows=11, cols= 4)
        tabla_top.style = 'Light Grid Accent 1'

        #encabezados
        encabezados = ["País", "Población 2024", "Cambio %", "Continente"]
        for i,encabezado in enumerate(encabezados):
            tabla_top.rows[0].cells[i].text = encabezado

        #datos
        for idx, dato in enumerate(top10.iter_rows(named=True)):
            tabla_top.rows[idx + 1].cells[0].text = dato['pais']
            tabla_top.rows[idx + 1].cells[1].text = f"{dato['poblacion_2024']:,}"
            tabla_top.rows[idx + 1].cells[2].text = f"{dato['cambio_porcentual']:,}"
            tabla_top.rows[idx + 1].cells[3].text = dato['continente']

        logger.info("   guardando archivo...")
        fecha_archivo = datetime.now().strftime("%Y-%m-%d_%H-%M")
        nombre_archivo = f"reporte_población_{fecha_archivo}.docx"
        ruta_completa = f"outputs/{nombre_archivo}"

        doc.save(ruta_completa)

        #verificar que se creó el archivo
        if not Path(ruta_completa).exists():
            logger.error(f"❌ el archivo no se creó en: {ruta_completa}")
            return None

        tamaño_kb = Path(ruta_completa).start().st_size / 1024

        logger.info("=" * 60)
        logger.info("REPORTE GENERADO EXITOSAMENTE")
        logger.info(f"  archivo: {nombre_archivo}")
        logger.info(f"  ubicación: {ruta_completa}")
        logger.info(f"  tamaño: {tamaño_kb:.2f} KB")
        logger.info("=" * 60)

        return ruta_completa

    except Exception as e:
        logger.error("=" * 60)
        logger.error("❌ ERROR CRÍTICO EN GENERACIÓN DE REPORTE")
        logger.error(f"     tipo: {type(e).__name__}")
        logger.error(f"     detalle: {str(e)}")
        logger.error("=" * 60)
        return None


